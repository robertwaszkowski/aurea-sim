import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt
import os
from rich.console import Console
from aureasim.ai_generator import humanize_name

console = Console()
def generate_charts(results_df, output_path, kpi_config=None):
    """Generates and saves KPI charts using Seaborn."""
    print("[Reporter] Generating charts...")
    sns.set_theme(style="whitegrid")

    # Humanize scenario names for display (do not mutate the original df)
    plot_df = results_df.copy()
    if 'Scenario' in plot_df.columns:
        plot_df['Scenario'] = plot_df['Scenario'].apply(
            lambda s: s.replace('_', ' ') if isinstance(s, str) else s
        )

    fig, axes = plt.subplots(1, 2, figsize=(15, 6))
    
    # 1. Cycle Time Chart
    sns.barplot(data=plot_df, x='Scenario', y='Avg_Cycle_Time_Days', ax=axes[0], hue='Scenario', palette="viridis", legend=False)
    axes[0].set_title('Average Process Cycle Time (Days)')
    axes[0].tick_params(axis='x', rotation=15)
    
    # 2. Case Cost Chart
    sns.barplot(data=plot_df, x='Scenario', y='Avg_Cost_Per_Case_PLN', ax=axes[1], hue='Scenario', palette="autumn", legend=False)
    axes[1].set_title('Average Cost per Case (PLN)')
    axes[1].tick_params(axis='x', rotation=15)
    
    for ax in axes:
        for container in ax.containers: 
            ax.bar_label(container, fmt='%.1f', padding=3)
            
    plt.tight_layout()
    plt.savefig(output_path, dpi=300)
    plt.close()
    print(f"  -> Saved chart to: {output_path}")

def generate_offline_executive_summary(results_df) -> str:
    """Generates a rule-based executive summary analyzing the simulation results."""
    # Find best scenarios
    best_time_row = results_df.loc[results_df['Avg_Cycle_Time_Days'].idxmin()]
    best_cost_row = results_df.loc[results_df['Avg_Cost_Per_Case_PLN'].idxmin()]
    
    # Calculate difference
    baseline_row = results_df[results_df['Scenario'].str.contains('Base|A_', case=False, na=False)]
    if baseline_row.empty:
        baseline_row = results_df.iloc[0]
    else:
        baseline_row = baseline_row.iloc[0]
        
    summary = []
    summary.append("# Executive Summary (Rule-Based Analysis)\n")
    summary.append("This report presents a comparative analysis of the simulated business process scenarios, evaluating performance across cycle times, process execution costs, and resource constraints.\n")
    
    summary.append("## Key Findings\n")
    
    time_imp_str = ""
    if best_time_row['Scenario'] != baseline_row['Scenario'] and baseline_row['Avg_Cycle_Time_Days'] > 0:
        pct_improvement = round((baseline_row['Avg_Cycle_Time_Days'] - best_time_row['Avg_Cycle_Time_Days']) / baseline_row['Avg_Cycle_Time_Days'] * 100, 1)
        time_imp_str = f" This represents a **{pct_improvement}% improvement** compared to the baseline scenario ({baseline_row['Avg_Cycle_Time_Days']} days)."
    summary.append(f"- **Optimal Cycle Time**: Scenario **{best_time_row['Scenario']}** achieved the lowest average cycle time of **{best_time_row['Avg_Cycle_Time_Days']} days**.{time_imp_str}\n")
        
    cost_sav_str = ""
    if best_cost_row['Scenario'] != baseline_row['Scenario']:
        cost_savings = round(baseline_row['Avg_Cost_Per_Case_PLN'] - best_cost_row['Avg_Cost_Per_Case_PLN'], 2)
        cost_sav_str = f" This yields savings of **{cost_savings} PLN per case** over the baseline scenario ({baseline_row['Avg_Cost_Per_Case_PLN']} PLN)."
    summary.append(f"- **Optimal Cost Efficiency**: Scenario **{best_cost_row['Scenario']}** resulted in the lowest average operating cost per case of **{best_cost_row['Avg_Cost_Per_Case_PLN']} PLN**.{cost_sav_str}\n")
        
    # Analyze wait times / bottlenecks
    wait_cols = [c for c in results_df.columns if c.startswith('Wait_Time_Hrs_')]
    if wait_cols:
        summary.append("## Resource & Bottleneck Analysis\n")
        summary.append("Analyzing the wait times across resource roles:\n")
        for col in wait_cols:
            role_name = col.replace('Wait_Time_Hrs_', '')
            max_wait_row = results_df.loc[results_df[col].idxmax()]
            summary.append(f"- **{role_name}**: Experienced the highest wait time in scenario **{max_wait_row['Scenario']}** (**{max_wait_row[col]} hours**).\n")
            
    summary.append("\n## Conclusion & Recommendations\n")
    if best_time_row['Scenario'] == best_cost_row['Scenario']:
        summary.append(f"Based on the KPIs, **{best_time_row['Scenario']}** is the clear recommendation, achieving the optimal balance of both time efficiency and cost savings.")
    else:
        summary.append(f"Choosing the optimal scenario depends on strategic priorities:\n")
        summary.append(f"- For **faster process completion**, deploy **{best_time_row['Scenario']}** ({best_time_row['Avg_Cycle_Time_Days']} days).\n")
        summary.append(f"- For **lowest operational costs**, deploy **{best_cost_row['Scenario']}** ({best_cost_row['Avg_Cost_Per_Case_PLN']} PLN).\n")
        
    return "\n".join(summary)

def generate_docx_report(report_settings, results_df, charts_path, output_docx_path, ai_summary_text=None, dfs_dict=None, base_config=None, scenarios_config=None):
    """Compiles a professional Word Document with text, a dynamic KPI table, and embedded charts."""
    doc = Document()
    import datetime
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # --- Header ---
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "AureaSim Evaluation Report"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # --- Title Block ---
    title = report_settings.get("title", "Automated Simulation Experiment Report")
    doc.add_heading(title, 0)
    
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # --- Footer ---
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts = datetime.datetime.now().strftime('%Y-%m-%d %H:%M')
    footer_para.add_run(f"Generated: {ts} | Page ")
    
    # Add dynamic page number field
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    description = report_settings.get("description", "Below are the results computed from the Prosimos execution engine based on the defined scenarios.")
    doc.add_paragraph(description)
    
    # --- Body ---
    if ai_summary_text:
        from aureasim.ai_generator import compress_citations
        ai_summary_text = compress_citations(ai_summary_text)
        
        import re
        def clean_markdown(text):
            # Remove bold/italic markdown markers while keeping the text
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            return text

        doc.add_heading("Executive Summary", level=1)
        for line in ai_summary_text.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith('### '):
                doc.add_heading(clean_markdown(line.replace('### ', '').strip()), level=4)
            elif line.startswith('## '):
                doc.add_heading(clean_markdown(line.replace('## ', '').strip()), level=3)
            elif line.startswith('# '):
                doc.add_heading(clean_markdown(line.replace('# ', '').strip()), level=2)
            elif line.startswith('- ') or line.startswith('* '):
                # lstrip removes the '-' or '*' and any leading spaces
                clean_line = clean_markdown(line.lstrip('-* ').strip())
                doc.add_paragraph(clean_line, style='List Bullet')
            else:
                doc.add_paragraph(clean_markdown(line))

    # Simulation Setup Parameters (Table 4 & Table 5)
    if base_config or scenarios_config:
        doc.add_heading("Simulation Parameters & Setup Matrix", level=1)
        doc.add_paragraph("This section outlines the baseline process parameter configurations and the experimental what-if scenario overrides simulated.")
        
        if base_config:
            try:
                profiles = base_config.get('resource_profiles', [])
                if profiles:
                    doc.add_heading("Table 4. Baseline Resource Setup", level=2)
                    r_cols = ["Resource Role", "Base Capacity", "Hourly Cost Rate (PLN/hr)", "Calendar"]
                    rtable = doc.add_table(rows=1, cols=len(r_cols))
                    rtable.style = 'Medium Shading 1 Accent 1'
                    
                    hdr_cells = rtable.rows[0].cells
                    for i, c in enumerate(r_cols):
                        hdr_cells[i].text = c
                        
                    for prof in profiles:
                        role_id = prof.get('id', prof.get('name', 'Unknown'))
                        for r_item in prof.get('resource_list', []):
                            row_cells = rtable.add_row().cells
                            row_cells[0].text = humanize_name(role_id)
                            row_cells[1].text = str(r_item.get('amount', 1))
                            try:
                                cost = float(r_item.get('cost_per_hour', 0.0))
                                row_cells[2].text = f"{cost:.2f}"
                            except Exception:
                                row_cells[2].text = str(r_item.get('cost_per_hour', 0.0))
                            row_cells[3].text = str(r_item.get('calendar', 'Standard_Working_Hours')).replace('_', ' ')
            except Exception as e:
                doc.add_paragraph(f"[WARNING] Could not parse baseline resource configurations: {e}")

            try:
                arrival_flow = base_config.get('arrival_time', {}).get('frequency', {})
                if arrival_flow:
                    events = arrival_flow.get('events', 1.0)
                    per_count = arrival_flow.get('per_count', 1.0)
                    per_unit = arrival_flow.get('per_unit', 'hour')
                    doc.add_paragraph(f"Baseline case arrival rate: {events} cases per {per_count} {per_unit}(s).")
            except Exception:
                pass

        if scenarios_config:
            try:
                if isinstance(scenarios_config, dict):
                    sc_list = scenarios_config.get('scenarios', [])
                else:
                    sc_list = scenarios_config
                    
                if sc_list:
                    doc.add_heading("Table 5. Scenario Parameter Overrides Matrix", level=2)
                    s_cols = ["Scenario Name", "Target Workload / Arrival Overrides", "Resource Overrides"]
                    stable = doc.add_table(rows=1, cols=len(s_cols))
                    stable.style = 'Medium Shading 1 Accent 1'
                    
                    hdr_cells = stable.rows[0].cells
                    for i, c in enumerate(s_cols):
                        hdr_cells[i].text = c
                        
                    for sc in sc_list:
                        s_name = sc.get('name', 'Unknown')
                        
                        # Format arrival rate override
                        arr_rate = sc.get('arrival_rate')
                        arr_str = "Default Baseline Rate"
                        if arr_rate is not None:
                            arr_str = f"Mean Interval: {arr_rate} seconds"
                        
                        # Format resource allocations overrides
                        res_allocs = sc.get('resource_allocations', {})
                        if res_allocs:
                            alloc_lines = [f"{humanize_name(k)}: {v}" for k, v in res_allocs.items()]
                            res_str = ", ".join(alloc_lines)
                        else:
                            res_str = "Default Capacity"
                            
                        row_cells = stable.add_row().cells
                        row_cells[0].text = humanize_name(s_name)
                        row_cells[1].text = arr_str
                        row_cells[2].text = res_str
            except Exception as e:
                doc.add_paragraph(f"[WARNING] Could not parse scenario parameter matrix: {e}")
    
    # Dynamic KPI Data Table
    doc.add_heading("Simulation Metrics & Deep Dive Data", level=1)
    
    doc.add_heading("Table 1. Impact of Scenarios on Main Process KPIs", level=2)
    scenario_col = "Scenario" if "Scenario" in results_df.columns else results_df.columns[0]
    scenarios = results_df[scenario_col].tolist()
    kpis = [c for c in results_df.columns if c != scenario_col]
    
    table = doc.add_table(rows=1, cols=len(scenarios) + 1)
    table.style = 'Medium Shading 1 Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "KPI"
    for i, sc in enumerate(scenarios):
        hdr_cells[i+1].text = str(sc).replace("_", " ")
        
    for kpi in kpis:
        row_cells = table.add_row().cells
        row_cells[0].text = str(kpi).replace("_", " ")
        for i, idx in enumerate(results_df.index):
            row_cells[i+1].text = str(results_df.loc[idx, kpi])
            
    if dfs_dict:
        for s_name, df_s in dfs_dict.items():
            doc.add_heading(f"Scenario Specifics: {s_name}", level=2)
            
            # Prepare Table 2 data (Top 5 resources by cost)
            resource_stats = df_s[df_s['Hourly_Rate'] > 0].groupby('resource').agg(
                Task_Count=('case_id', 'count'), 
                Total_Work_Hrs=('work_time_hr', 'sum'),
                Total_Cost=('Task_Cost_PLN', 'sum')
            ).sort_values(by='Total_Cost', ascending=False).head(5).round(2).reset_index()
            
            doc.add_heading(f"Table 2. Aggregated Resource Statistics (Top 5) - {s_name}", level=3)
            if not resource_stats.empty:
                r_cols = ["Resource", "Executions", "Total Work Time (h)", "Total Cost (PLN)"]
                rtable = doc.add_table(rows=1, cols=len(r_cols))
                rtable.style = 'Medium Shading 1 Accent 1'
                r_cells = rtable.rows[0].cells
                for i, c in enumerate(r_cols):
                    r_cells[i].text = c
                
                for _, r_row in resource_stats.iterrows():
                    row_cells = rtable.add_row().cells
                    row_cells[0].text = humanize_name(str(r_row['resource']))
                    row_cells[1].text = str(r_row['Task_Count'])
                    row_cells[2].text = str(r_row['Total_Work_Hrs'])
                    row_cells[3].text = str(r_row['Total_Cost'])
            else:
                doc.add_paragraph("No resources with specific hourly rates found for this scenario.")
                
            # Prepare Table 3 data (Top 5 activities by cost)
            activity_stats = df_s.groupby(['activity', 'resource']).agg(
                Task_Count=('case_id', 'count'),
                Avg_Work_Hrs=('work_time_hr', 'mean'),
                Total_Cost=('Task_Cost_PLN', 'sum')
            ).sort_values(by='Total_Cost', ascending=False).head(5).round(2).reset_index()
            
            doc.add_heading(f"Table 3. Key Process Activities by Cost and Volume - {s_name}", level=3)
            if not activity_stats.empty:
                a_cols = ["Task Name", "Performer", "Executions", "Avg Task Time (h)", "Total Cost (PLN)"]
                atable = doc.add_table(rows=1, cols=len(a_cols))
                atable.style = 'Medium Shading 1 Accent 1'
                a_cells = atable.rows[0].cells
                for i, c in enumerate(a_cols):
                    a_cells[i].text = c
                
                for _, a in activity_stats.iterrows():
                    cells = atable.add_row().cells
                    cells[0].text = str(a['activity'])
                    cells[1].text = humanize_name(str(a['resource']))
                    cells[2].text = str(a['Task_Count'])
                    cells[3].text = str(a['Avg_Work_Hrs'])
                    cells[4].text = str(a['Total_Cost'])
            else:
                doc.add_paragraph("No activity cost data found for this scenario.")
        
    # Embed Charts safely
    if os.path.exists(charts_path):
        doc.add_heading("Performance Visualizations", level=1)
        doc.add_picture(charts_path, width=Inches(6.0))
        
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("The simulation has completed successfully across all defined scenarios. Costing and cycle times were evaluated using Activity-Based Costing (ABC) principles with dynamic staff scaling schedules.")

    # References section from metadata source_urls
    base_meta = {}
    if dfs_dict and hasattr(dfs_dict, 'get'):
        pass  # dfs_dict is scenario data, not metadata
    
    all_refs = report_settings.get("references", []) if isinstance(report_settings, dict) else []

    if all_refs:
        doc.add_heading("Sources", level=1)
        for i, ref in enumerate(all_refs, 1):
            title = ref.get("title", ref.get("url", ""))
            url = ref.get("url", "")
            p = doc.add_paragraph(style='List Number')
            if url.startswith("http"):
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                # The 'List Number' style auto-numbers, so skip manual "[i]" prefix
                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('r:id'),
                    p.part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True))
                new_run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                rStyle = OxmlElement('w:rStyle')
                rStyle.set(qn('w:val'), 'Hyperlink')
                rPr.append(rStyle)
                new_run.append(rPr)
                t = OxmlElement('w:t')
                t.text = title
                new_run.append(t)
                hyperlink.append(new_run)
                p._p.append(hyperlink)
            else:
                p.add_run(f"{title}")

    doc.save(output_docx_path)
    console.print(f"  [green]✔  DOCX report generated:[/green] {os.path.basename(output_docx_path)}")

def generate_pdf_report(report_settings, results_df, charts_path, output_pdf_path, ai_summary_text=None, dfs_dict=None, base_config=None, scenarios_config=None):
    """Compiles a PDF document using fpdf2 with text, tables, and embedded charts."""
    try:
        from fpdf import FPDF
    except ImportError:
        console.print("  [red]fpdf2 not installed. Please run: pip install fpdf2[/red]")
        return
        
    class PDFReport(FPDF):
        """Use a bundled Unicode font even where legacy calls ask for Helvetica."""
        def set_font(self, family=None, style="", size=0):
            if family and family.casefold() == "helvetica":
                family = "AureaUnicode"
            return super().set_font(family, style, size)

        def header(self):
            self.set_font("helvetica", "B", 10)
            self.set_text_color(0, 188, 212) # Cyan
            self.cell(0, 10, "AureaSim Evaluation Report", ln=True, align="R")
            self.set_text_color(0, 0, 0)
            
        def footer(self):
            self.set_y(-15)
            self.set_font("helvetica", "I", 8)
            self.set_text_color(128, 128, 128)
            import datetime
            ts = datetime.datetime.now().strftime('%Y-%m-%d %H:%M')
            self.cell(0, 10, f"Generated: {ts} | Page {self.page_no()}", align="C")
            self.set_text_color(0, 0, 0)

    pdf = PDFReport()
    # Core FPDF fonts such as Helvetica only support Latin-1.  DejaVu Sans is
    # distributed with Matplotlib (already required for AureaSim charts) and
    # supports Polish task and process names.
    font_dir = os.path.join(matplotlib.get_data_path(), "fonts", "ttf")
    for style, filename in {
        "": "DejaVuSans.ttf",
        "B": "DejaVuSans-Bold.ttf",
        "I": "DejaVuSans-Oblique.ttf",
        "BI": "DejaVuSans-BoldOblique.ttf",
    }.items():
        pdf.add_font("AureaUnicode", style=style, fname=os.path.join(font_dir, filename))
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # Title
    title = report_settings.get("title", "Automated Simulation Experiment Report")
    pdf.set_font("helvetica", "B", 18)
    pdf.multi_cell(0, 10, title, align="C")
    pdf.ln(5)
    
    # Description
    description = report_settings.get("description", "Below are the results computed from the Prosimos execution engine based on the defined scenarios.")
    pdf.set_font("helvetica", "", 11)
    pdf.multi_cell(0, 7, description)
    pdf.ln(5)
    
    if ai_summary_text:
        pdf.set_font("helvetica", "", 10)
        for line in ai_summary_text.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            import re
            # Prevent unbreakable character sequences that crash FPDF wrapper
            line = re.sub(r'[-_*=]{15,}', '----------', line)
            
            # FPDF2 crashes when tabs are rendered
            line = line.replace('\t', '    ')
            
            # Replace non-breaking spaces which Gemini often produces. They prevent word-wrapping.
            line = line.replace(u'\xa0', u' ')
            
            # Prevent single extremely long unbroken words (like filepaths) that crash FPDF 
            words = line.split(' ')
            line = ' '.join([w if len(w) < 40 else ' '.join(w[i:i+40] for i in range(0, len(w), 40)) for w in words])
            
            line = line.replace('**', '') # Strip bolding, fpdf text styling can be tricky inline without markdown support.
            
            try:
                # Force left margin cursor to unconditionally prevent negative w_max
                pdf.set_x(pdf.l_margin)
                # Use strict explicit effective page width instead of 0
                w = pdf.epw 
                
                if line.startswith('### '):
                    pdf.set_font("helvetica", "B", 12)
                    pdf.multi_cell(w, 8, line.replace('### ', ''))
                    pdf.set_font("helvetica", "", 10)
                elif line.startswith('## '):
                    pdf.set_font("helvetica", "B", 14)
                    pdf.multi_cell(w, 8, line.replace('## ', ''))
                    pdf.set_font("helvetica", "", 10)
                elif line.startswith('# '):
                    pdf.set_font("helvetica", "B", 16)
                    pdf.multi_cell(w, 10, line.replace('# ', ''))
                    pdf.set_font("helvetica", "", 10)
                elif line.startswith('- ') or line.startswith('* '):
                    # Use the Unicode bullet supported by the embedded DejaVu font.
                    # chr(149) is a Windows-1252 control character, not a Unicode bullet.
                    pdf.multi_cell(w, 6, "\u2022 " + line[2:])
                else:
                    pdf.multi_cell(w, 6, line)
            except Exception as e:
                pass # Silently skip unrenderable text lines to prevent console spam
                
        pdf.ln(5)
            
    pdf.set_font("helvetica", "B", 14)
    pdf.cell(0, 10, "Simulation Metrics & Deep Dive Data", ln=True)
    pdf.ln(2)
    
    # Basic Main KPI Table
    pdf.set_font("helvetica", "B", 12)
    pdf.cell(0, 8, "Table 1. Impact of Scenarios on Main Process KPIs", ln=True)
    
    # Very simple table using generic cells
    scenario_col = "Scenario" if "Scenario" in results_df.columns else results_df.columns[0]
    scenarios = results_df[scenario_col].tolist()
    kpis = [c for c in results_df.columns if c != scenario_col]
    
    pivoted_cols = ["KPI"] + scenarios
    col_width = pdf.epw / max(1, len(pivoted_cols))
    
    # Table Header
    pdf.set_fill_color(0, 188, 212) # Cyan header
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("helvetica", "B", 9)
    for col in pivoted_cols:
        col_text = str(col).replace("_", " ")
        if len(col_text) > 25: col_text = col_text[:22] + "..."
        pdf.cell(col_width, 8, col_text, border=1, align="C", fill=True)
    pdf.ln()
    pdf.set_text_color(0, 0, 0)
    
    # Table Rows
    pdf.set_font("helvetica", "", 9)
    fill = False
    for kpi in kpis:
        if fill:
            pdf.set_fill_color(240, 240, 240)
        else:
            pdf.set_fill_color(255, 255, 255)
        
        kpi_text = str(kpi).replace("_", " ")
        if len(kpi_text) > 25: kpi_text = kpi_text[:22] + "..."
        pdf.cell(col_width, 8, kpi_text, border=1, align="L", fill=True)
        
        for idx in results_df.index:
            val = str(results_df.loc[idx, kpi])
            if len(val) > 25: val = val[:22] + "..."
            pdf.cell(col_width, 8, val, border=1, align="C", fill=True)
        pdf.ln()
        fill = not fill
    pdf.ln(5)

    if dfs_dict:
        for s_name, df_s in dfs_dict.items():
            pdf.add_page()
            pdf.set_font("helvetica", "B", 14)
            pdf.cell(0, 10, f"Scenario Specifics: {s_name}", ln=True)
            pdf.ln(3)

            # Table 2: Resource Stats
            resource_stats = df_s[df_s['Hourly_Rate'] > 0].groupby('resource').agg(
                Task_Count=('case_id', 'count'), 
                Total_Work_Hrs=('work_time_hr', 'sum'),
                Total_Cost=('Task_Cost_PLN', 'sum')
            ).sort_values(by='Total_Cost', ascending=False).head(5).round(2).reset_index()

            pdf.set_font("helvetica", "B", 12)
            pdf.cell(0, 8, f"Table 2. Aggregated Resource Statistics (Top 5) - {s_name}", ln=True)
            
            if not resource_stats.empty:
                r_cols = ["Resource", "Executions", "Total Work Time (h)", "Total Cost (PLN)"]
                rt_width = pdf.epw / 4
                pdf.set_fill_color(0, 188, 212) # Cyan header
                pdf.set_text_color(255, 255, 255)
                pdf.set_font("helvetica", "B", 9)
                for c in r_cols:
                    pdf.cell(rt_width, 8, c, border=1, align="C", fill=True)
                pdf.ln()
                pdf.set_text_color(0, 0, 0)
                pdf.set_font("helvetica", "", 9)
                fill = False
                for _, r in resource_stats.iterrows():
                    if fill:
                        pdf.set_fill_color(240, 240, 240)
                    else:
                        pdf.set_fill_color(255, 255, 255)
                    pdf.cell(rt_width, 8, str(r['resource']), border=1, fill=True)
                    pdf.cell(rt_width, 8, str(r['Task_Count']), border=1, align="C", fill=True)
                    pdf.cell(rt_width, 8, str(r['Total_Work_Hrs']), border=1, align="C", fill=True)
                    pdf.cell(rt_width, 8, str(r['Total_Cost']), border=1, align="C", fill=True)
                    pdf.ln()
                    fill = not fill
            else:
                pdf.set_font("helvetica", "I", 10)
                pdf.cell(0, 8, "No resource cost data found.", ln=True)
            pdf.ln(5)

            # Table 3: Activity Stats
            activity_stats = df_s.groupby(['activity', 'resource']).agg(
                Task_Count=('case_id', 'count'),
                Avg_Work_Hrs=('work_time_hr', 'mean'),
                Total_Cost=('Task_Cost_PLN', 'sum')
            ).sort_values(by='Total_Cost', ascending=False).head(5).round(2).reset_index()

            pdf.set_font("helvetica", "B", 12)
            pdf.cell(0, 8, f"Table 3. Key Process Activities by Cost and Volume - {s_name}", ln=True)
            
            if not activity_stats.empty:
                a_cols = ["Task Name", "Performer", "Executions", "Avg Time (h)", "Cost (PLN)"]
                at_width = pdf.epw / 5
                pdf.set_fill_color(0, 188, 212) # Cyan header
                pdf.set_text_color(255, 255, 255)
                pdf.set_font("helvetica", "B", 9)
                for c in a_cols:
                    pdf.cell(at_width, 8, c, border=1, align="C", fill=True)
                pdf.ln()
                pdf.set_text_color(0, 0, 0)
                pdf.set_font("helvetica", "", 8)
                fill = False
                for _, a in activity_stats.iterrows():
                    if fill:
                        pdf.set_fill_color(240, 240, 240)
                    else:
                        pdf.set_fill_color(255, 255, 255)
                    pdf.cell(at_width, 8, str(a['activity'])[:20], border=1, fill=True)
                    pdf.cell(at_width, 8, humanize_name(str(a['resource'])), border=1, fill=True)
                    pdf.cell(at_width, 8, str(a['Task_Count']), border=1, align="C", fill=True)
                    pdf.cell(at_width, 8, str(a['Avg_Work_Hrs']), border=1, align="C", fill=True)
                    pdf.cell(at_width, 8, str(a['Total_Cost']), border=1, align="C", fill=True)
                    pdf.ln()
                    fill = not fill
            else:
                pdf.set_font("helvetica", "I", 10)
                pdf.cell(0, 8, "No activity cost data found.", ln=True)
    
    # Images
    if os.path.exists(charts_path):
        pdf.add_page()
        pdf.set_font("helvetica", "B", 14)
        pdf.cell(0, 10, "Performance Visualizations", ln=True)
        pdf.ln(5)
        # Assuming typical landscape chart from seaborn
        img_w = pdf.epw # Use full page width
        pdf.image(charts_path, w=img_w)
        
    pdf.ln(10)
    pdf.set_font("helvetica", "B", 14)
    pdf.cell(0, 10, "Conclusion", ln=True)
    pdf.set_font("helvetica", "", 11)
    pdf.multi_cell(0, 7, "The simulation has completed successfully across all defined scenarios. Costing and cycle times were evaluated using Activity-Based Costing (ABC) principles with dynamic staff scaling schedules.")

    # References section from references list
    refs = report_settings.get("references", []) if isinstance(report_settings, dict) else []
    if refs:
        pdf.add_page()
        pdf.set_font("helvetica", "B", 14)
        pdf.cell(0, 10, "Sources", ln=True)
        pdf.ln(3)
        pdf.set_font("helvetica", "", 9)
        for i, ref in enumerate(refs, 1):
            title = ref.get("title", ref.get("url", ""))
            display = title
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(pdf.epw, 6, f"[{i}] {display}")
            pdf.ln(1)

    try:
        pdf.output(output_pdf_path)
        console.print(f"  [green][OK] PDF report generated:[/green] {os.path.basename(output_pdf_path)}")
    except Exception as e:
        console.print(f"  [red][ERROR] PDF report generation failed:[/red] {e}")
        raise RuntimeError(f"PDF report generation failed: {e}") from e


def escape_latex(text: str) -> str:
    """Escapes all LaTeX special characters safely."""
    if not isinstance(text, str):
        text = str(text)
    # The order is critical: backslash must be escaped first
    text = text.replace('\\', '\\textbackslash{}')
    # Then escape standard characters
    replacements = {
        '&': '\\&', '%': '\\%', '$': '\\$', '#': '\\#', '_': '\\_',
        '{': '\\{', '}': '\\}',
        '~': '\\textasciitilde{}', '^': '\\textasciicircum{}'
    }
    for char, rep in replacements.items():
        text = text.replace(char, rep)
    return text

def generate_latex_report(report_settings, results_df, charts_path, output_tex_path, ai_summary_text=None, dfs_dict=None, base_config=None, scenarios_config=None):
    """Generates a raw LaTeX source file for the report."""
    
    title = escape_latex(report_settings.get("title", "Automated Simulation Experiment Report"))
    description = escape_latex(report_settings.get("description", "Results from Prosimos simulation."))
    
    tex_content = []
    tex_content.append("\\documentclass[12pt, a4paper]{article}")
    tex_content.append("\\usepackage[utf8]{inputenc}")
    tex_content.append("\\usepackage{graphicx}")
    tex_content.append("\\usepackage{booktabs}")
    tex_content.append("\\usepackage{geometry}")
    tex_content.append("\\usepackage{longtable}")
    tex_content.append("\\usepackage{tabularx}")
    tex_content.append("\\usepackage{xcolor}")
    tex_content.append("\\usepackage{tcolorbox}")
    tex_content.append("\\tcbuselibrary{breakable}")
    tex_content.append("\\usepackage{float}")
    tex_content.append("\\usepackage[hidelinks]{hyperref}")
    tex_content.append("\\usepackage{fancyhdr}")
    tex_content.append("\\geometry{a4paper, margin=1in}")
    tex_content.append("\\pagestyle{fancy}")
    tex_content.append("\\fancyhf{}")
    import datetime
    ts = datetime.datetime.now().strftime('%Y-%m-%d %H:%M')
    tex_content.append(f"\\cfoot{{\\color{{gray}} Generated: {ts} | Page \\thepage}}")
    tex_content.append("\\renewcommand{\\headrulewidth}{0pt}")
    tex_content.append("")
    tex_content.append(f"\\title{{{title}}}")
    tex_content.append("\\author{AureaSim Engine}")
    tex_content.append("\\date{\\today}")
    tex_content.append("")
    tex_content.append("\\begin{document}")
    tex_content.append("\\sloppy")
    tex_content.append("\\maketitle")
    tex_content.append("\\thispagestyle{fancy}")
    tex_content.append("")
    tex_content.append(f"\\noindent {description}")
    tex_content.append("\\vspace{1em}")
    tex_content.append("")
    
    if ai_summary_text:
        from aureasim.ai_generator import compress_citations
        ai_summary_text = compress_citations(ai_summary_text)
        tex_content.append("\\section*{Executive Summary}")
        tex_content.append("\\begin{tcolorbox}[colback=blue!5!white,colframe=blue!75!black,title=AI Generated Insights,breakable]")
        for line in ai_summary_text.split('\n'):
            line = line.strip()
            if not line:
                continue
            # Basic bold markdown to tex conversion
            if '**' in line:
                import re
                line = re.sub(r'\*\*(.*?)\*\*', r'\\textbf{\1}', line)
                
            if line.startswith('### '):
                content = escape_latex(line.replace('### ', ''))
                tex_content.append(f"\\textbf{{{content}}}\\\\[0.5em]")
            elif line.startswith('## '):
                content = escape_latex(line.replace('## ', ''))
                tex_content.append(f"\\textbf{{{content}}}\\\\[0.5em]")
            elif line.startswith('# '):
                content = escape_latex(line.replace('# ', ''))
                tex_content.append(f"\\textbf{{{content}}}\\\\[0.5em]")
            elif line.startswith('- ') or line.startswith('* '):
                content = escape_latex(line[2:])
                tex_content.append(f"$\\bullet$ {content}\\\\")
            else:
                line = escape_latex(line)
                tex_content.append(line + "\\\\[0.3em]")
        tex_content.append("\\end{tcolorbox}")
        tex_content.append("\\vspace{1em}")
                
    tex_content.append("\\section{Simulation Metrics}")
    tex_content.append("Table 1: Impact of Scenarios on Main Process KPIs")
    tex_content.append("")
    
    # Table
    scenario_col = "Scenario" if "Scenario" in results_df.columns else results_df.columns[0]
    scenarios = results_df[scenario_col].tolist()
    kpis = [c for c in results_df.columns if c != scenario_col]
    
    pivoted_cols = ["KPI"] + scenarios
    # Use X for KPI (left-aligned wrap) and centered X for scenarios (centered wrap). Remove vertical bars for booktabs.
    col_format = "X " + " ".join([">{\\centering\\arraybackslash}X"] * len(scenarios))
    
    tex_content.append("\\begin{table}[H]")
    tex_content.append("\\centering")
    tex_content.append(f"\\begin{{tabularx}}{{\\textwidth}}{{{col_format}}}")
    tex_content.append("\\toprule")
    
    headers = " & ".join([escape_latex(str(c).replace('_', ' ')) for c in pivoted_cols])
    formatted_headers = headers.replace(' & ', '} & \\textbf{')
    tex_content.append(f"\\textbf{{{formatted_headers}}} \\\\")
    tex_content.append("\\midrule")
    
    for kpi in kpis:
        kpi_clean = escape_latex(str(kpi).replace('_', ' '))
        r_vals = [kpi_clean]
        for idx in results_df.index:
            val = escape_latex(str(results_df.loc[idx, kpi]))
            r_vals.append(val)
        tex_content.append(" & ".join(r_vals) + " \\\\")
        
    tex_content.append("\\bottomrule")
    tex_content.append("\\end{tabularx}")
    tex_content.append("\\end{table}")
    tex_content.append("")
    
    if dfs_dict:
        for s_name, df_s in dfs_dict.items():
            tex_name = escape_latex(s_name)
            tex_content.append(f"\\subsection*{{Scenario Specifics: {tex_name}}}")

            # Data 2: Resource Stats
            resource_stats = df_s[df_s['Hourly_Rate'] > 0].groupby('resource').agg(
                Task_Count=('case_id', 'count'), 
                Total_Work_Hrs=('work_time_hr', 'sum'),
                Total_Cost=('Task_Cost_PLN', 'sum')
            ).sort_values(by='Total_Cost', ascending=False).head(5).round(2).reset_index()

            tex_content.append(f"\\subsubsection*{{Table 2. Aggregated Resource Statistics (Top 5) - {tex_name}}}")
            if not resource_stats.empty:
                tex_content.append("\\begin{table}[H]")
                tex_content.append("\\begin{tabularx}{\\textwidth}{X l r r}")
                tex_content.append("\\toprule")
                tex_content.append("\\textbf{Resource} & \\textbf{Executions} & \\textbf{Work Time (h)} & \\textbf{Cost (PLN)} \\\\")
                tex_content.append("\\midrule")
                for _, r in resource_stats.iterrows():
                    res_name = escape_latex(str(r['resource']))
                    tex_content.append(f"{res_name} & {r['Task_Count']} & {r['Total_Work_Hrs']} & {r['Total_Cost']} \\\\")
                tex_content.append("\\bottomrule")
                tex_content.append("\\end{tabularx}")
                tex_content.append("\\end{table}")
                tex_content.append("\\vspace{1em}")

            # Data 3: Activity Stats
            activity_stats = df_s.groupby(['activity', 'resource']).agg(
                Task_Count=('case_id', 'count'),
                Avg_Work_Hrs=('work_time_hr', 'mean'),
                Total_Cost=('Task_Cost_PLN', 'sum')
            ).sort_values(by='Total_Cost', ascending=False).head(5).round(2).reset_index()

            tex_content.append(f"\\subsubsection*{{Table 3. Key Process Activities by Cost and Volume - {tex_name}}}")
            if not activity_stats.empty:
                tex_content.append("\\begin{table}[H]")
                tex_content.append("\\begin{tabularx}{\\textwidth}{X X r r r}")
                tex_content.append("\\toprule")
                tex_content.append("\\textbf{Task} & \\textbf{Performer} & \\textbf{Exec.} & \\textbf{Avg Time (h)} & \\textbf{Cost (PLN)} \\\\")
                tex_content.append("\\midrule")
                for _, a in activity_stats.iterrows():
                    act_name = escape_latex(str(a['activity']))
                    res_name2 = escape_latex(humanize_name(str(a['resource'])))
                    tex_content.append(f"{act_name} & {res_name2} & {a['Task_Count']} & {a['Avg_Work_Hrs']} & {a['Total_Cost']} \\\\")
                tex_content.append("\\bottomrule")
                tex_content.append("\\end{tabularx}")
                tex_content.append("\\end{table}")
                tex_content.append("\\vspace{1.5em}")
        
    if os.path.exists(charts_path):
        import pathlib
        p = pathlib.Path(charts_path).absolute().as_posix() # Needs forward slashes for TeX
        tex_content.append("\\section{Performance Visualizations}")
        tex_content.append("\\begin{figure}[h!]")
        tex_content.append("\\centering")
        tex_content.append(f"\\includegraphics[width=\\textwidth]{{{p}}}")
        tex_content.append("\\caption{Scenario Comparison Charts}")
        tex_content.append("\\end{figure}")
        
    tex_content.append("\\section{Conclusion}")
    tex_content.append("The simulation has completed successfully across all defined scenarios. Costing and cycle times were evaluated using Activity-Based Costing (ABC) principles with dynamic staff scaling schedules.")
    tex_content.append("")

    # References section from references list
    refs = report_settings.get("references", []) if isinstance(report_settings, dict) else []
    if refs:
        tex_content.append("\\section{Sources}")
        tex_content.append("\\begin{enumerate}")
        for ref in refs:
            title = ref.get("title", "")
            url = ref.get("url", "")
            safe_title = title.replace('&', '\\&').replace('%', '\\%').replace('#', '\\#').replace('_', '\\_')
            if url.startswith("http"):
                tex_content.append(f"  \\item \\href{{{url}}}{{{safe_title}}}")
            else:
                tex_content.append(f"  \\item {safe_title}")
        tex_content.append("\\end{enumerate}")
        tex_content.append("")

    tex_content.append("\\end{document}")
    
    with open(output_tex_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(tex_content))
        
    console.print(f"  [green]✔  LaTeX source generated:[/green] {os.path.basename(output_tex_path)}")
